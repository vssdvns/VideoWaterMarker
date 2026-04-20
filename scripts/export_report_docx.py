from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEX_PATH = ROOT / "docs" / "MASTER_FINAL_REPORT.tex"
OUT_PATH = ROOT / "MASTER_FINAL_REPORT.docx"

REPORT_TITLE = (
    "HYBRID VIDEO WATERMARKING FOR OTT PIRACY PREVENTION USING CONTENT-AWARE "
    "PLACEMENT, FORENSIC TRACEABILITY, AND ROBUST DETECTION"
)
STUDENT_NAME = "SubhaSai DurgaVenkata Vinnakota"
DEGREE_NAME = "MASTER OF SCIENCE"
MAJOR_NAME = "Computer Science"
DEPARTMENT_NAME = "Computer Science"
TERM_NAME = "SPRING 2026"
COMMITTEE_CHAIR = "Dr. Bang Tran S"
SECOND_READER = "Dr. Ying Jin"
GRAD_COORDINATOR = "Dr. Haiquan Chen"


def set_page_number_footer(section, start: int | None = None) -> None:
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    if start is not None:
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)
        pg_num_type.set(qn("w:start"), str(start))


def configure_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1.5)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        set_page_number_footer(section, start=1)


def add_centered(doc: Document, text: str, *, bold: bool = False, size: int = 12, spacing_after: int = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)


def add_body(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 3)}"
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12 if level > 1 else 14)
    if level == 1:
        run.bold = True


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def latex_to_text(text: str) -> str:
    replacements = {
        r"\ReportTitle": REPORT_TITLE,
        r"\StudentName": STUDENT_NAME,
        r"\DegreeName": DEGREE_NAME,
        r"\MajorName": MAJOR_NAME,
        r"\DepartmentName": DEPARTMENT_NAME,
        r"\TermName": "SPRING",
        r"\TermYear": "2026",
        r"\CommitteeChair": COMMITTEE_CHAIR,
        r"\SecondReader": SECOND_READER,
        r"\GraduateCoordinator": GRAD_COORDINATOR,
        r"\%": "%",
        r"\alpha": "alpha",
        r"\texttt{": "",
        r"\textbf{": "",
        r"\textit{": "",
        r"\emph{": "",
        r"\par": " ",
        r"\\": " ",
        r"~": " ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\$\\?[\^_]\{([^}]*)\}\$", r"^\1", text)
    text = text.replace("$", "")
    text = re.sub(r"\\[A-Za-z]+\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[A-Za-z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_body(tex: str) -> list[tuple[str, str]]:
    start = tex.index(r"\chapter{INTRODUCTION}")
    end = tex.index(r"\begin{thebibliography}{99}")
    body = tex[start:end]
    lines = body.splitlines()
    items: list[tuple[str, str]] = []
    para: list[str] = []
    in_figure = False
    in_table = False
    pending_caption = ""
    pending_alt = ""

    def flush_paragraph() -> None:
        nonlocal para
        if para:
            text = latex_to_text(" ".join(para))
            if text:
                items.append(("paragraph", text))
        para = []

    for raw in lines:
        line = raw.strip()
        if not line:
            flush_paragraph()
            continue
        if line.startswith(r"\begin{figure"):
            flush_paragraph()
            in_figure = True
            pending_caption = ""
            pending_alt = ""
            continue
        if line.startswith(r"\end{figure"):
            if pending_caption:
                items.append(("figure", latex_to_text(pending_caption)))
            if pending_alt:
                items.append(("alt", latex_to_text(pending_alt)))
            in_figure = False
            continue
        if line.startswith(r"\begin{table"):
            flush_paragraph()
            in_table = True
            pending_caption = ""
            pending_alt = ""
            continue
        if line.startswith(r"\end{table"):
            if pending_caption:
                items.append(("table", latex_to_text(pending_caption)))
            if pending_alt:
                items.append(("alt", latex_to_text(pending_alt)))
            in_table = False
            continue
        if in_figure or in_table:
            if line.startswith(r"\caption{"):
                pending_caption = line[len(r"\caption{") : -1]
            elif line.startswith(r"\AltText{"):
                pending_alt = line[len(r"\AltText{") : -1]
            continue
        if line.startswith(r"\chapter{"):
            flush_paragraph()
            items.append(("chapter", latex_to_text(line[len(r"\chapter{") : -1])))
            continue
        if line.startswith(r"\section{"):
            flush_paragraph()
            items.append(("section", latex_to_text(line[len(r"\section{") : -1])))
            continue
        if line.startswith(r"\subsection{"):
            flush_paragraph()
            items.append(("subsection", latex_to_text(line[len(r"\subsection{") : -1])))
            continue
        if line.startswith(r"\clearpage") or line.startswith(r"\begingroup") or line.startswith(r"\endgroup"):
            flush_paragraph()
            continue
        if line.startswith("\\"):
            continue
        para.append(line)

    flush_paragraph()
    return items


def extract_references(tex: str) -> list[str]:
    refs_block = tex.split(r"\begin{thebibliography}{99}", 1)[1].split(r"\end{thebibliography}", 1)[0]
    refs = re.findall(r"\\bibitem\{[^}]+\}\s*(.*?)(?=(?:\\bibitem\{|$))", refs_block, flags=re.S)
    return [latex_to_text(ref) for ref in refs if latex_to_text(ref)]


def build_front_matter(doc: Document) -> None:
    add_centered(doc, REPORT_TITLE, bold=True, size=14, spacing_after=18)
    add_centered(doc, "A Project", size=12, spacing_after=18)
    add_centered(doc, f"Presented to the faculty of the Department of {DEPARTMENT_NAME}", size=12)
    add_centered(doc, "California State University, Sacramento", size=12, spacing_after=18)
    add_centered(doc, "Submitted in partial satisfaction of", size=12)
    add_centered(doc, "the requirements for the degree of", size=12, spacing_after=12)
    add_centered(doc, DEGREE_NAME, size=12)
    add_centered(doc, "in", size=12)
    add_centered(doc, MAJOR_NAME, size=12, spacing_after=18)
    add_centered(doc, "by", size=12)
    add_centered(doc, STUDENT_NAME, size=12, spacing_after=18)
    add_centered(doc, "SPRING", size=12)
    add_centered(doc, "2026", size=12)
    add_page_break(doc)

    add_centered(doc, REPORT_TITLE, size=12, spacing_after=18)
    add_centered(doc, "A Project", size=12, spacing_after=18)
    add_centered(doc, "by", size=12)
    add_centered(doc, STUDENT_NAME, size=12, spacing_after=18)
    add_body(doc, "Approved by:")
    add_body(doc, "________________________________________")
    add_body(doc, f"{COMMITTEE_CHAIR}, Committee Chair")
    add_body(doc, "Date")
    add_body(doc, "________________________________________")
    add_body(doc, f"{SECOND_READER}, Second Reader")
    add_body(doc, "Date")
    add_page_break(doc)

    add_body(doc, f"Student: {STUDENT_NAME}")
    add_body(
        doc,
        "I certify that this student has met the requirements for format contained in the University format manual, "
        "and this project is suitable for electronic submission to the library and credit is to be awarded for the project.",
    )
    add_body(doc, "________________________________________    Date")
    add_body(doc, f"{GRAD_COORDINATOR}, Graduate Coordinator")
    add_body(doc, f"Department of {DEPARTMENT_NAME}")
    add_page_break(doc)

    add_centered(doc, "Abstract", size=12)
    add_centered(doc, "of", size=12)
    add_centered(doc, REPORT_TITLE, size=12)
    add_centered(doc, "by", size=12)
    add_centered(doc, STUDENT_NAME, size=12, spacing_after=12)


def main() -> None:
    tex = TEX_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    build_front_matter(doc)

    abstract_problem = (
        "This project studies how to watermark video in a way that is practical for over-the-top streaming platforms. "
        "The main goal is to place a visible watermark where it is less distracting while still keeping it recoverable "
        "after common distortions such as re-encoding, blur, crop, grayscale conversion, frame-rate reduction, and resize. "
        "A second goal is traceability, so that a leaked copy can still be tied back to a user or session. To address both "
        "needs, the project combines visible watermarking, invisible payload embedding, detection, attack testing, and a "
        "user-facing application in one workflow."
    )
    abstract_data = (
        "The implemented system uses Laplacian-based complexity, semantic saliency from pretrained models, temporal smoothing, "
        "optional optical flow, visible text rendering, discrete cosine transform payload embedding, and a neural watermarking branch. "
        "It also stores frame-level placement metadata in a positions.json file so that detection can use both local search and a "
        "global fallback when the watermark shifts after attack. The broader project workflow draws on a standardized 150-clip "
        "training corpus for model development and a separate 26-clip HD benchmark for visible robustness evaluation."
    )
    abstract_conclusion = (
        "The main conclusion is that content-aware visible watermarking is the strongest and most mature part of the system. "
        "The DCT and neural branches make the design more complete and more useful for forensic tracing, but they still involve "
        "stronger quality-versus-robustness tradeoffs."
    )
    add_body(doc, "Statement of Problem")
    add_body(doc, abstract_problem)
    add_body(doc, "Sources of Data")
    add_body(doc, abstract_data)
    add_body(doc, "Conclusions Reached")
    add_body(doc, abstract_conclusion)
    add_body(doc, "________________________________________")
    add_body(doc, f"{COMMITTEE_CHAIR}, Committee Chair")
    add_body(doc, "Date")
    add_page_break(doc)

    add_heading(doc, "ACKNOWLEDGEMENTS", 1)
    add_body(
        doc,
        "I would like to express my sincere gratitude to my project advisor for the guidance, patience, and insights that "
        "shaped every stage of this work. The regular feedback sessions were essential in refining both the software and this written report.",
    )
    add_body(
        doc,
        "I am grateful to the faculty of the Department of Computer Science at California State University, Sacramento, "
        "whose coursework in software engineering, data science, and related areas provided an important foundation for this project. "
        "I also thank the second reader and committee members for suggestions that improved the clarity of the manuscript.",
    )
    add_body(
        doc,
        "Special thanks go to the open-source communities behind OpenCV, PyTorch, Streamlit, FastAPI, and FFmpeg, whose freely available "
        "tools made this project possible. Finally, I thank my family and friends for their encouragement and support throughout my graduate studies.",
    )
    add_page_break(doc)

    for kind, text in extract_body(tex):
        if kind == "chapter":
            add_heading(doc, text, 1)
        elif kind == "section":
            add_heading(doc, text, 2)
        elif kind == "subsection":
            add_heading(doc, text, 3)
        elif kind == "paragraph":
            add_body(doc, text)
        elif kind == "figure":
            add_body(doc, f"Figure: {text}")
        elif kind == "table":
            add_body(doc, f"Table: {text}")
        elif kind == "alt":
            add_body(doc, f"Alternative text: {text}")

    add_page_break(doc)
    add_heading(doc, "REFERENCES", 1)
    for idx, ref in enumerate(extract_references(tex), start=1):
        add_body(doc, f"{idx}. {ref}")

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
